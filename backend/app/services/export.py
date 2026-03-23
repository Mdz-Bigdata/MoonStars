"""
文档导出服务 (PDF, Word)
"""
import io
import os
import re
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Inches, Pt
from xhtml2pdf import pisa
from uuid import UUID

import logging
logger = logging.getLogger(__name__)

class ExportService:
    def __init__(self):
        pass
        
    def _clean_html(self, html: str) -> str:
        """移除简单的 HTML 标签以获取纯文本"""
        if not html:
            return ""
        return re.sub(r'<[^>]+>', '', str(html))

    async def to_word(self, title: str, content: List[Dict[str, Any]]) -> io.BytesIO:
        """导出为 Word"""
        doc = Document()
        doc.add_heading(title, 0)
        
        for block in content:
            b_type = block.get('type')
            b_data = block.get('content', {}) or {}
            
            if b_type == 'text':
                p = doc.add_paragraph(self._clean_html(b_data.get('text', '')))
            elif b_type == 'heading':
                level = b_data.get('level', 2)
                doc.add_heading(self._clean_html(b_data.get('text', '')), level=level-1 if level > 1 else 1)
            elif b_type == 'image':
                # 注意：实际生产中需要下载图片并处理
                doc.add_paragraph(f"[图片: {b_data.get('url')}]")
            elif b_type == 'list':
                for item in b_data.get('items', []):
                    item_text = item.get('text', '') if isinstance(item, dict) else str(item)
                    doc.add_paragraph(self._clean_html(item_text), style='List Bullet')
            elif b_type == 'table':
                rows = b_data.get('rows', [])
                if rows:
                    num_cols = len(rows[0].get('cells', []))
                    table = doc.add_table(rows=len(rows), cols=num_cols)
                    for r_idx, row in enumerate(rows):
                        for c_idx, cell in enumerate(row.get('cells', [])):
                            cell_text = self._clean_html(cell.get('content', [{}])[0].get('content', {}).get('text', ''))
                            table.cell(r_idx, c_idx).text = cell_text
                            
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

    async def to_pdf(self, title: str, content: List[Dict[str, Any]]) -> io.BytesIO:
        """导出为 PDF (通过 HTML 中转)"""
        html_content = f"<html><head><meta charset='utf-8'><title>{title}</title><style>body {{ font-family: 'Helvetica'; }} h1 {{ text-align: center; }}</style></head><body>"
        html_content += f"<h1>{title}</h1>"
        
        for block in content:
            b_type = block.get('type')
            b_data = block.get('content', {}) or {}
            
            if b_type == 'text':
                html_content += f"<p>{b_data.get('text', '')}</p>"
            elif b_type == 'heading':
                level = b_data.get('level', 2)
                html_content += f"<h{level}>{b_data.get('text', '')}</h{level}>"
            elif b_type == 'image':
                html_content += f"<div style='text-align:center'><img src='{b_data.get('url')}' style='max-width:100%'></div>"
            elif b_type == 'list':
                html_content += "<ul>"
                for item in b_data.get('items', []):
                    item_text = item.get('text', '') if isinstance(item, dict) else str(item)
                    html_content += f"<li>{item_text}</li>"
                html_content += "</ul>"
            elif b_type == 'table':
                html_content += "<table border='1' style='width:100%; border-collapse: collapse;'>"
                for row in b_data.get('rows', []):
                    html_content += "<tr>"
                    for cell in row.get('cells', []):
                        cell_text = cell.get('content', [{}])[0].get('content', {}).get('text', '')
                        html_content += f"<td>{cell_text}</td>"
                    html_content += "</tr>"
                html_content += "</table>"
                
        html_content += "</body></html>"
        
        file_stream = io.BytesIO()
        pisa_status = pisa.CreatePDF(html_content, dest=file_stream)
        
        if pisa_status.err:
            logger.error("PDF 转换失败")
            raise Exception("PDF 转换失败")
            
        file_stream.seek(0)
        return file_stream

    async def export_purchase_records_to_pdf(self, username: str, records: List[Any]) -> io.BytesIO:
        """从购买记录导出 PDF"""
        html_content = f"<html><head><meta charset='utf-8'><title>{username} 的购买记录</title>"
        html_content += "<style>body { font-family: 'Helvetica'; } table { width:100%; border-collapse: collapse; } th, td { border: 1px solid #ddd; padding: 8px; text-align: left; } th { background-color: #f2f2f2; }</style></head><body>"
        html_content += f"<h1>{username} 的购买记录</h1>"
        html_content += "<table><thead><tr><th>时间</th><th>商名</th><th>金额</th><th>状态</th></tr></thead><tbody>"
        
        for r in records:
            time_str = r.created_at.strftime("%Y-%m-%d %H:%M")
            column_name = r.column.name if r.column else "未知专栏"
            status_map = {"paid": "已支付", "pending": "待支付", "failed": "支付失败", "cancelled": "已取消", "shipped": "已发货", "completed": "已完成"}
            status_text = status_map.get(r.status.value if hasattr(r.status, "value") else r.status, str(r.status))
            
            html_content += f"<tr><td>{time_str}</td><td>{column_name}</td><td>{r.amount / 100:.2f} 元</td><td>{status_text}</td></tr>"
            
        html_content += "</tbody></table></body></html>"
        
        file_stream = io.BytesIO()
        pisa.CreatePDF(html_content, dest=file_stream)
        file_stream.seek(0)
        return file_stream
